from collections import defaultdict
from collections import Counter
from re import S
import pandas as pd
import os
import json
import logging
import basic_preprocessing
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from basic_preprocessing import BasicPreprocessor
from config import *

bp = BasicPreprocessor()


class TopTokens:
    def __init__(self, input_dir, output_dir = "outputs"):
        self.input_dir = input_dir
        self.intent_each_token_count = {}
        self.intents = []
        self.vocabulary = []
        self.total_terminals = None
        self.output_dir = output_dir


    def all_CSVs_to_df(self, extension = '.csv'): # Non default argument dump_dir must be before default arguments
        dfs = []
        try:
            for subdir, dirs, files in os.walk(self.input_dir): # os.walk() requires folder not file
                for filename in files:
                    ext = os.path.splitext(filename)[-1].lower()
                    if ext == extension:
                        input_file_dir = os.path.join(subdir, filename)
                        try:
                            df = pd.read_csv(input_file_dir, header = None)
                        except:
                            logging.error("Cannot read knowledgebase for language detection")

                        dfs.append(pd.DataFrame({"sample_text": df.iloc[:, 0], "intent": filename}))
            return pd.concat(dfs)
        except Exception as e:
            logging.error(e)


    def tokenizer(self, example, ngram = TOKENIZATION_RANGE):
        if len(ngram) == 1:
            ngram.append(ngram[0])
        ngram = sorted(ngram)
        
        assert len(ngram) >=1 and len(ngram) <= 2, "The ngram parameter must be list of length 1 or 2."

        splitted_example = str(example).split()
        tokens = []
        for ng in range(ngram[0], ngram[1] + 1, 1):
            for i in range(len(splitted_example) - ng + 1):
                tokens.append(' '.join(splitted_example[i: i + ng]))
        return tokens


    def get_vocabulary_counter(self):
        if self.vocabulary == []:
            df = self.all_CSVs_to_df()
            examples = list(df["sample_text"])
            _terminals = []

            for example in examples:
                example = str(example)
                example = bp.do_basic_preprocessing(text = example)
                _terminals.extend(self.tokenizer(example))
            self.total_terminals = len(_terminals)
            _terminals = list(set(_terminals))
            _terminals.sort()
            self.vocabulary = _terminals

        dd = defaultdict(int, { k: 0 for k in self.vocabulary})

        return dd


    def intentwise_each_token_count(self):
        if self.intent_each_token_count == {}:
            intentwise_token_count = {}
            df = self.all_CSVs_to_df()
            if self.intents == []:
                intents = list(set(df["intent"]))
                intents.sort()
                self.intents = intents
            
            for intent in self.intents:
                examples_under_intent = df[df["intent"] == intent]["sample_text"]
                vb_counter = self.get_vocabulary_counter()

                for example in examples_under_intent:
                    example = str(example)
                    example = bp.do_basic_preprocessing(text = example)
                    
                    for terminal in self.tokenizer(example):
                        vb_counter[terminal] += 1

                intentwise_token_count[intent] = vb_counter
            
            self.intent_each_token_count = intentwise_token_count

        return self.intent_each_token_count


    def token_probability_per_intent(self): # Denominator
        itc = self.intentwise_each_token_count()
        token_av_per_int = {}
        for intent in self.intents:
            # print(intent)
            tokens_with_cnt = itc[intent]
            # print(tokens_with_cnt)
            cum = 0
            for k, v in tokens_with_cnt.items():
                cum += int(v)
            token_av_per_int[intent] = cum / self.total_terminals
        # with open("tapi.txt", "w") as f:
        #     f.write(str(token_av_per_int))
        return token_av_per_int


    def probability_of_token_given_intent(self, take_positive_probabilities_only = True):
        iwtc = self.intentwise_each_token_count()
        iwtc_div_total_tokens = {k: {_k: _v / self.total_terminals for _k, _v in v.items()} for k, v in iwtc.items()}
        tppi = self.token_probability_per_intent() # Denominator

        for intent in self.intents:
            numerator = iwtc_div_total_tokens[intent]
            denominator = tppi[intent]
            proba = {k: v / denominator for k, v in numerator.items()}
            if take_positive_probabilities_only:
                proba = {k: v for k, v in sorted(proba.items(), key=lambda item: item[1], reverse = True) if v > 0}
            else:
                proba = {k: v for k, v in sorted(proba.items(), key=lambda item: item[1], reverse = True)}

            if not os.path.isdir(f"{self.output_dir}/intentwise_token_probabilities"):
                os.makedirs(f"{self.output_dir}/intentwise_token_probabilities")
            with open(f"{self.output_dir}/intentwise_token_probabilities/{''.join(intent.split('.')[:-1])}.json", "w") as jf:
                json.dump(proba, jf, ensure_ascii=False)
            # with open(f"{self.output_dir}/intentwise_token_probabilities/{''.join(intent.split('.')[:-1])}.txt", "w") as f:
            #     f.write(str(proba))


    def remove_all_dupes(self, d):
        c = Counter()
        for v in d.values():
            c.update(v)

        for k,v in d.items():
            d[k] = [item for item in v if c[item] == 1]
        return d


    def find_unique_tokens(self, 
        take_minimum_identifiers = True,
        input_file_extensions = [".json"]): #TODO: take_minimum_identifier: if some token is unique in unigram, no need to include it in any bigram
        
        intentwise_tokens = {}
        for subdir, dirs, files in os.walk(f"{self.output_dir}/intentwise_token_probabilities"):
            for file in files:
                ext = os.path.splitext(file)[-1].lower()
                if ext in input_file_extensions:
                    # print(file)
                    input_file = os.path.join(subdir, file)
                    print(input_file)
                    with open(input_file) as jf:
                        d = json.load(jf)
                    intentwise_tokens[file] = list(d.keys())
        uniques = self.remove_all_dupes(intentwise_tokens)
        
        if not os.path.isdir(f"{self.output_dir}/intentwise_unique_tokens"):
            os.makedirs(f"{self.output_dir}/intentwise_unique_tokens")
        for file in uniques.keys():
            d = uniques[file]

            if take_minimum_identifiers:
                """If some token is unique in some lower ngram, we need not take it in association 
                with higher ngrams and we can discard all higher ngrams. For example, assuming we
                are working with bigram and some unique tokens found as: 'token1', 'token1 token2a',
                'token1 token2b', 'token1 token2c'.
                Here, 'token1' is enough to distinguish this intent, so we don't need:
                'token1 token2a', 'token1 token2b', 'token1 token2c' etc. and we can discard them all."""
                d.sort(key = lambda x: len(x.split()), reverse = False) # sorting based on token count (descending)
                checkerboard = set()
                _d = []
                
                for token in d:
                    splitted_token = token.split()
                    take_it = not any(element in checkerboard for element in splitted_token)
                    
                    if take_it:
                        _d.append(token)
                        checkerboard.update(splitted_token)
                
                _d.sort() # Now it is to time to sort according to lexicographical order for better readability
                d = _d

            with open(f"{self.output_dir}/intentwise_unique_tokens/{file}", "w") as jf:
                json.dump(d, jf, ensure_ascii=False)


    def generate_data_not_containing_top_tokens(self):
        df = self.all_CSVs_to_df()
        if self.intents == []:
                intents = list(set(df["intent"]))
                intents.sort()
                self.intents = intents

        if not os.path.isdir("quarantined"):
            os.makedirs("quarantined")
            
        for intent in self.intents:
            examples_under_intent = df[df["intent"] == intent]["sample_text"]
            with open(f"outputs/{intent}.json", "r") as jf:
                tokens = json.load(jf)
            keys = tokens.keys()
            quarantined = []
            for example in examples_under_intent:
                splitted = self.tokenizer(example)
                take_it = False
                for term in splitted:
                    if term in keys:
                        take_it = True
                        break
                if take_it == False:
                    quarantined.append(example)
            pd.DataFrame(quarantined).to_csv(f"quarantined/{intent}.csv", index = False)

    


    
    def generate_annotated_output(self, apply_min_max_normalization = True):
        basic_preprocessing.load_stopwords(force_loading = True)

        if basic_preprocessing.predefined_stopwords:
            pass
        else:
            logging.error("No predefined stopwords found. Please make sure there is a csv file containing stopwords.")
        
        df = self.all_CSVs_to_df()
        intents = set(df["intent"])

        def change_range(old_value, old_mini = 0, old_max = 1, new_mini = 0, new_max = 255):
            return int(((old_value - old_mini) / (old_max - old_mini)) * (new_max - new_mini) + new_mini)
        
        for intent in intents:
            intended_examples = df[df["intent"] == intent]["sample_text"]

            with open(f"outputs/intentwise_token_probabilities/{'.'.join(intent.split('.')[:-1])}.json") as jf:
                intended_token_probabilities = json.load(jf)

            if apply_min_max_normalization:
                # For better visualization, as the probabilities are too small for each of
                # the tokens, hence, it is hard to distinguish them while viewing if we do not normalize them.
                probabilities = intended_token_probabilities.values()
                mini = min(probabilities)
                maxm = max(probabilities)
                intended_token_probabilities = {k: (v - mini) / (maxm - mini) for k, v in intended_token_probabilities.items()}

            doc = docx.Document()
            para = doc.add_paragraph('')

            for example in intended_examples.iloc:
                example = bp.do_basic_preprocessing(text = example, remove_predefined_stopwords = False)
                # Not considering bigram and higher grams for now
                splitted_example = str(example).split()
                for terminal in splitted_example:
                    if terminal in intended_token_probabilities.keys():
                        if terminal in basic_preprocessing.predefined_stopwords:
                            # print(change_range(intended_token_probabilities[terminal]))
                            para.add_run(f" {terminal}").font.color.rgb = RGBColor(change_range(intended_token_probabilities[terminal]), 0, 0)
                            # para.add_run(f" {terminal}").font.color.rgb = RGBColor(255, 0, 0)
                        else:
                            para.add_run(f" {terminal}").font.color.rgb = RGBColor(0, 0, change_range(intended_token_probabilities[terminal]))
                            # para.add_run(f" {terminal}").font.color.rgb = RGBColor(0, 0, 255)
                    else:
                        print(terminal)
                        para.add_run(f" {terminal}").font.color.rgb = RGBColor(211, 211, 211)
                para.add_run('\n')

            if not os.path.isdir(f"{self.output_dir}/annotated_docs"):
                os.makedirs(f"{self.output_dir}/annotated_docs")
            doc.save(f"{self.output_dir}/annotated_docs/{'.'.join(intent.split('.')[:-1])}.docx")

        basic_preprocessing.predefined_stopwords = None # Unloading it again as it was loaded as needed
        


if __name__ == "__main__":
    tt = TopTokens(input_dir = "intent-sample")
    # print(tt.get_vocabulary_counter()[:100])
    # print(tt.get_vocabulary_counter()[:100])
    tt.intentwise_each_token_count()
    # tt.probability_of_token_given_intent(take_positive_probabilities_only = False)
    tt.probability_of_token_given_intent(take_positive_probabilities_only = True)
    # tt.generate_data_not_containing_top_tokens()
    tt.find_unique_tokens()
    
    tt.generate_annotated_output()